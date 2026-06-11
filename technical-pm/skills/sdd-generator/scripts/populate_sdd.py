
#!/usr/bin/env python3
"""Populate the bundled SDD template from a sdd_content.json file."""

import json
import sys
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

INSTRUCTION_PARA_MARKERS = [
    'List the key people involved in the building of the agent',
    'List all agent technical dependencies - the underlying systems',
    'Visually outline the sequence of steps the agent follows from start to finish',
    'List all the prompts used in the agent. Use links',
    'Describe how each failure scenario is addressed',
    'Outline the mandatory tests required to validate',
]


def clear_instruction_para(para):
    """Remove all runs from a paragraph, leaving it empty but structurally intact."""
    for run in para.runs:
        run.text = ''
    # Also clear any raw XML text nodes
    for child in para._p:
        if child.tag == qn('w:r'):
            for t in child.findall(qn('w:t')):
                t.text = ''


def set_para_text(para, text):
    """Set paragraph text, preserving the formatting of the first run."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ''


def set_cell(cell, text):
    """Set the text of the first paragraph in a table cell."""
    if cell.paragraphs:
        set_para_text(cell.paragraphs[0], str(text) if text else '')
    else:
        cell.add_paragraph(str(text) if text else '')


def ensure_rows(table, n_data_rows):
    """Ensure table has at least header_row + n_data_rows rows, cloning the last row as needed."""
    while len(table.rows) < n_data_rows + 1:
        # Clone the last row
        last_row = table.rows[-1]._tr
        new_row = deepcopy(last_row)
        table._tbl.append(new_row)


def populate_sdd(content_path, template_path, output_path):
    with open(content_path, encoding='utf-8') as f:
        c = json.load(f)

    doc = Document(template_path)

    # Clear instruction-only paragraphs
    for para in doc.paragraphs:
        for marker in INSTRUCTION_PARA_MARKERS:
            if marker in para.text:
                clear_instruction_para(para)
                break

    tables = doc.tables

    # ── Table 0: Document History ──────────────────────────────────────────────
    if len(tables) > 0:
        t = tables[0]
        if len(t.rows) > 1:
            row = t.rows[1]
            cells = row.cells
            if len(cells) >= 4:
                set_cell(cells[0], c.get('version', '1.0'))
                set_cell(cells[1], c.get('date', ''))
                set_cell(cells[2], c.get('author', '[TBD]'))
                set_cell(cells[3], 'Initial draft')

    # ── Purpose paragraph ──────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if 'Purpose' in para.text and para.style.name.startswith('Heading'):
            # Find the next non-empty, non-heading paragraph after this heading
            found_heading = False
            for p in doc.paragraphs:
                if p == para:
                    found_heading = True
                    continue
                if found_heading and p.style.name.startswith('Heading'):
                    break
                if found_heading and p.text.strip():
                    set_para_text(p, c.get('purpose', ''))
                    break
            break

    # ── Table 1: Objectives ────────────────────────────────────────────────────
    if len(tables) > 1:
        t = tables[1]
        obj = c.get('objectives', {})
        field_map = {
            'Problem': obj.get('problem', '[TBD]'),
            'Objective': obj.get('objective', '[TBD]'),
            'Value': obj.get('value', '[TBD]'),
            'Expected outcome': obj.get('expected_outcome', '[TBD]'),
        }
        for row in t.rows:
            cells = row.cells
            if len(cells) >= 2:
                label = cells[0].text.strip()
                for key, val in field_map.items():
                    if key.lower() in label.lower():
                        set_cell(cells[1], val)
                        break

    # ── Table 2: Key Contacts ──────────────────────────────────────────────────
    if len(tables) > 2:
        t = tables[2]
        contacts = c.get('key_contacts', [])
        ensure_rows(t, len(contacts))
        for i, contact in enumerate(contacts):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 3:
                set_cell(cells[0], contact.get('role', ''))
                set_cell(cells[1], contact.get('name', '[TBD]'))
                set_cell(cells[2], contact.get('email', '[TBD]'))

    # ── Process as-is and high-level solution paragraphs ──────────────────────
    # These are typically free-text paragraphs found by proximity to their headings
    _fill_section_para(doc, 'as-is', c.get('process_as_is', ''))
    _fill_section_para(doc, 'high-level solution', c.get('high_level_solution', ''))
    _fill_section_para(doc, 'high level solution', c.get('high_level_solution', ''))

    # ── Table 3: Applications ──────────────────────────────────────────────────
    if len(tables) > 3:
        t = tables[3]
        apps = c.get('applications', [])
        ensure_rows(t, len(apps))
        for i, app in enumerate(apps):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                set_cell(cells[0], app.get('name', ''))
            if len(cells) >= 2:
                set_cell(cells[1], app.get('purpose', ''))
            if len(cells) >= 3:
                set_cell(cells[2], app.get('owner', '[TBD]'))
            if len(cells) >= 4:
                set_cell(cells[3], app.get('access_method', '[TBD]'))

    # ── Table 4: Process Flow ──────────────────────────────────────────────────
    if len(tables) > 4:
        t = tables[4]
        flow = c.get('process_flow', [])
        ensure_rows(t, len(flow))
        for i, step in enumerate(flow):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                set_cell(cells[0], step.get('step', str(i + 1)))
            if len(cells) >= 2:
                set_cell(cells[1], step.get('action', ''))
            if len(cells) >= 3:
                set_cell(cells[2], step.get('actor', ''))
            if len(cells) >= 4:
                set_cell(cells[3], step.get('output', ''))

    # ── Process steps narrative ────────────────────────────────────────────────
    _fill_section_para(doc, 'process steps', c.get('process_steps_detail', ''))
    _fill_section_para(doc, 'step detail', c.get('process_steps_detail', ''))

    # ── Table 5: Restrictions ──────────────────────────────────────────────────
    if len(tables) > 5:
        t = tables[5]
        restrictions = c.get('restrictions', [])
        ensure_rows(t, len(restrictions))
        for i, r in enumerate(restrictions):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                set_cell(cells[0], r.get('restriction', ''))
            if len(cells) >= 2:
                set_cell(cells[1], r.get('rationale', ''))

    # ── Prompt logic paragraph ─────────────────────────────────────────────────
    _fill_section_para(doc, 'prompt logic', c.get('prompt_logic', ''))
    _fill_section_para(doc, 'prompting', c.get('prompt_logic', ''))

    # ── Table 6: Prompt list ───────────────────────────────────────────────────
    if len(tables) > 6:
        t = tables[6]
        prompts = c.get('prompts', [])
        ensure_rows(t, len(prompts))
        for i, p in enumerate(prompts):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                set_cell(cells[0], p.get('name', ''))
            if len(cells) >= 2:
                set_cell(cells[1], p.get('purpose', ''))
            if len(cells) >= 3:
                set_cell(cells[2], p.get('link', '[TBD]'))

    # ── Table 7: Fallback scenarios ────────────────────────────────────────────
    FALLBACK_KEYS = [
        'no_data_found',
        'api_timeout',
        'low_confidence',
        'user_out_of_scope',
        'auth_failure',
        'duplicate_detected',
        'data_quality_issue',
        'escalation_needed',
        'scheduled_job_failure',
    ]
    FALLBACK_LABELS = [
        'No data found',
        'API timeout',
        'Low confidence',
        'Out of scope',
        'Auth failure',
        'Duplicate detected',
        'Data quality issue',
        'Escalation needed',
        'Scheduled job failure',
    ]
    if len(tables) > 7:
        t = tables[7]
        fb = c.get('fallback_actions', {})
        ensure_rows(t, len(FALLBACK_KEYS))
        for i, (key, label) in enumerate(zip(FALLBACK_KEYS, FALLBACK_LABELS)):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                # Only overwrite scenario label if cell looks empty or generic
                existing = cells[0].text.strip()
                if not existing or existing.lower() in ('scenario', 'failure scenario'):
                    set_cell(cells[0], label)
            if len(cells) >= 2:
                set_cell(cells[1], fb.get(key, '[TBD]'))

    # ── Table 8: Test cases ────────────────────────────────────────────────────
    if len(tables) > 8:
        t = tables[8]
        tests = c.get('test_cases', [])
        ensure_rows(t, len(tests))
        for i, test in enumerate(tests):
            row_idx = i + 1
            if row_idx >= len(t.rows):
                break
            row = t.rows[row_idx]
            cells = row.cells
            if len(cells) >= 1:
                set_cell(cells[0], test.get('id', f'TC-{i+1:02d}'))
            if len(cells) >= 2:
                set_cell(cells[1], test.get('scenario', ''))
            if len(cells) >= 3:
                set_cell(cells[2], test.get('expected', ''))

    # ── Tail narrative sections ────────────────────────────────────────────────
    _fill_section_para(doc, 'challenge', c.get('challenges', ''))
    _fill_section_para(doc, 'deployment contact', c.get('deployment_contacts', ''))
    _fill_section_para(doc, 'deployment strategy', c.get('deployment_strategy', ''))
    _fill_section_para(doc, 'infrastructure', c.get('infrastructure', ''))
    _fill_section_para(doc, 'training', c.get('training', ''))
    _fill_section_para(doc, 'feedback', c.get('feedback_loop', ''))

    doc.save(output_path)
    print(f"Saved: {output_path}")


def _fill_section_para(doc, keyword, text):
    """Find the first content paragraph after a heading containing keyword, set its text."""
    if not text:
        return
    found_heading = False
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and keyword.lower() in para.text.lower():
            found_heading = True
            continue
        if found_heading:
            if para.style.name.startswith('Heading'):
                break
            if para.text.strip() and not any(m in para.text for m in INSTRUCTION_PARA_MARKERS):
                set_para_text(para, text)
                return
    # If no existing paragraph found, this section may be empty — skip silently


if __name__ == '__main__':
    if len(sys.argv) != 4:
        print("Usage: populate_sdd.py <content.json> <template.docx> <output.docx>")
        sys.exit(1)
    populate_sdd(sys.argv[1], sys.argv[2], sys.argv[3])
